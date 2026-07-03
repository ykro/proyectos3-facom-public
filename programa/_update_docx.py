"""Actualiza el .docx EN SITIO: semana 6 = login y datos por usuario, semana 7 =
cámara y pulido de marca. Preserva estructura, estilos, anchos y celdas fusionadas."""
import copy
from docx import Document
from docx.oxml.ns import qn

DOC = "/Users/ykro/code/temp-projects/proyectos3-programa-update/programa/programa-proyectos-iii-2026.docx"

d = Document(DOC)

def first_rpr(paragraph):
    if paragraph.runs:
        rpr = paragraph.runs[0]._r.find(qn('w:rPr'))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None

def add_formatted(paragraph, text, rpr):
    lines = text.split('\n')
    for i, ln in enumerate(lines):
        run = paragraph.add_run(ln)
        if rpr is not None:
            run._r.insert(0, copy.deepcopy(rpr))
        if i < len(lines) - 1:
            run.add_break()

def set_cell_text(cell, text):
    paras = cell.paragraphs
    first = paras[0]
    rpr = first_rpr(first)
    for r in list(first.runs):
        r._r.getparent().remove(r._r)
    for p in paras[1:]:
        p._p.getparent().remove(p._p)
    add_formatted(first, text, rpr)

cron = d.tables[2]

def set_row(r, fecha, tema, alumno, docente):
    cells = cron.rows[r].cells
    set_cell_text(cells[0], fecha)
    set_cell_text(cells[1], tema)
    set_cell_text(cells[2], alumno)
    set_cell_text(cells[3], docente)

# Semana 6: login y datos por usuario
set_row(6, "Semana 6 (10 de agosto)",
    "Desarrollo nativo II: login y datos por usuario",
    "Laboratorio 5: App nativa con login con Google y datos privados por usuario en Firestore/Storage",
    "Exposición sobre:\n- Inicio de sesión con Google (Firebase Authentication)\n"
    "- Datos privados por usuario en Firestore (uid, tiempo real)\n"
    "- Fotos en Firebase Storage\n- Depuración de UI con el modo de anotación y capturas de la vista previa")

# Semana 7: cámara y pulido de marca
set_row(7, "Semana 7 (17 de agosto)",
    "Desarrollo nativo III: cámara y pulido de marca",
    "Investigación 1: Reglas de seguridad de Firebase\n"
    "Laboratorio 6: App nativa con cámara, ícono/splash y lista para entregar",
    "Exposición sobre:\n- La cámara del teléfono y permisos en tiempo de ejecución\n"
    "- Reglas de seguridad de Firestore y Storage\n- Ícono, splash y pulido de marca\n"
    "- Depurar con evidencia: capturas de la vista previa y modo de anotación")

d.save(DOC)
print("Actualizado:", DOC)
