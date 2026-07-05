"""Edita el .docx original cambiando SOLO el texto, preservando estructura,
anchos de columna, celdas de examen fusionadas y estilos."""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = "/Users/ykro/code/cursos/proyectos3-facom/programa/_original-2024.docx"
OUT = "/Users/ykro/code/cursos/proyectos3-facom/programa/programa-proyectos-iii-2026.docx"

d = Document(SRC)

def first_rpr(paragraph):
    if paragraph.runs:
        rpr = paragraph.runs[0]._r.find(qn('w:rPr'))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None

def add_formatted(paragraph, text, rpr):
    lines = text.split('\n')
    for i, ln in enumerate(lines):
        run = paragraph.add_run(ln)
        if rpr is not None:
            run._r.insert(0, copy.deepcopy(rpr))
        if i < len(lines) - 1:
            run.add_break()

def set_para_text(paragraph, text):
    rpr = first_rpr(paragraph)
    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)
    add_formatted(paragraph, text, rpr)

def set_cell_text(cell, text):
    paras = cell.paragraphs
    first = paras[0]
    rpr = first_rpr(first)
    for r in list(first.runs):
        r._r.getparent().remove(r._r)
    for p in paras[1:]:
        p._p.getparent().remove(p._p)
    add_formatted(first, text, rpr)

def clone_after(ref_para, text):
    new = copy.deepcopy(ref_para._p)
    ref_para._p.addnext(new)
    np = Paragraph(new, ref_para._parent)
    set_para_text(np, text)
    return np

P = d.paragraphs

# --- Presentación (body) ---
set_para_text(P[4],
    "Este curso enseña a diseñar y construir aplicaciones móviles reales mediante el "
    "desarrollo de software asistido por inteligencia artificial. El estudiante de "
    "Comunicación y Diseño aprende a llevar una idea desde el concepto visual hasta una "
    "aplicación funcional y publicable, sin necesidad de dominar la programación "
    "tradicional. Usando Google AI Studio —que construye tanto aplicaciones web instalables "
    "(PWA) como aplicaciones nativas de Android, todo desde el navegador—, el estudiante "
    "describe en español lo que quiere construir y la IA genera "
    "el código, lo que le permite concentrarse en su fortaleza profesional: la experiencia "
    "de usuario, la identidad visual y la comunicación de producto.")

# --- Duración (body) ---
set_para_text(P[8],
    "El curso está diseñado para ser impartido en 10 sesiones de una hora y media cada una, "
    "en modalidad en línea y en dos secciones con el mismo contenido (Sección 1 los lunes y "
    "Sección 2 los miércoles). Se impartirán en 48 horas totales de trabajo del estudiante, "
    "incluidas las 15 horas académicas en sesiones de una hora y media cada una, con "
    "acompañamiento docente y las 33 horas restantes se emplearán en actividades "
    "independientes de estudio, prácticas, proyectos, preparación de exámenes u otras que "
    "sean necesarias para alcanzar las metas de aprendizaje propuestas, sin incluir las "
    "destinadas a la presentación de exámenes finales.")

# --- Competencia general (body) ---
set_para_text(P[12],
    "Diseñar y construir aplicaciones móviles funcionales (web instalables y nativas) "
    "mediante desarrollo asistido por inteligencia artificial, aplicando principios de "
    "experiencia de usuario, identidad visual y comunicación de producto.")

# --- Competencias específicas (lista: 2 existentes + 3 nuevas) ---
set_para_text(P[16],
    "Aplicar principios de UI/UX y diseño mobile-first en la creación de aplicaciones móviles.")
set_para_text(P[17],
    "Construir aplicaciones funcionales dirigiendo a la inteligencia artificial en "
    "Google AI Studio, sin escribir código manualmente.")
last = P[17]
for t in [
    "Integrar datos, autenticación de usuarios y capacidades del dispositivo en una "
    "aplicación, comprendiendo qué hace cada pieza.",
    "Distribuir y comunicar un producto digital: publicación, marketing básico, pitch y "
    "portafolio profesional.",
    "Aumentar la calidad académica del estudiante a través de la promoción de virtudes y "
    "valores humanos (flexibilidad, audacia, innovación).",
]:
    last = clone_after(last, t)

# --- Valor académico (body) ---
set_para_text(P[22],
    "Estimular la destreza del estudiante para convertir una idea en un producto digital "
    "real, combinando criterio de diseño con el uso responsable y efectivo de herramientas "
    "de inteligencia artificial, de forma que adquiera un diferenciador concreto para su "
    "vida profesional.")

# --- Metodología (body: original + taller IA) ---
set_para_text(P[26],
    "Se imparte a través de experiencias de aprendizaje en espacios simulados y/o "
    "controlados de prácticas, organizados al interior de la institución académica, "
    "destinados a formar en las habilidades técnicas, estratégicas, creativas y de "
    "indagación para la resolución de problemas reales de la práctica profesional y "
    "artística. El curso se desarrolla como un taller en línea asistido por IA: cada sesión "
    "combina la exposición de conceptos con laboratorios prácticos donde el estudiante "
    "construye, sobre dos proyectos integradores que evolucionan a lo largo del trimestre.")

# --- Actividad del estudiante (5 items existentes) ---
acts = [
    "Crear contenidos teóricos que respalden las prácticas del diseño de aplicaciones móviles.",
    "Realización de ejercicios creativos y laboratorios para la aplicación de cada tema, "
    "construyendo aplicaciones con inteligencia artificial.",
    "Proyectos de aplicación que acerquen al estudiante a una realidad profesional: una "
    "aplicación web instalable (PWA) y una aplicación nativa.",
    "Documentar su proceso mediante una bitácora de prompts y un caso de portafolio.",
    "Participación en talleres, ejercicios grupales y otras actividades de diseño para "
    "poner en práctica los conocimientos adquiridos.",
]
for idx, t in zip([31, 32, 33, 34, 35], acts):
    set_para_text(P[idx], t)

# --- Observaciones (2 existentes + 2 nuevas) ---
last = P[44]
for t in [
    "Los exámenes Parcial y Final son 100 % escritos y en línea (teoría + práctica), sin "
    "defensa: los 100 puntos de cada instrumento se convierten íntegramente a los 20 puntos "
    "del examen. Los proyectos (PWA y app nativa + portafolio + pitch) se entregan en línea "
    "y se evalúan dentro de Proyectos; el Demo Day de la sesión 10 es un showcase de "
    "celebración y entrega, no una defensa calificada.",
    "Los 30 puntos de Proyectos corresponden a los dos proyectos (PWA y app nativa + "
    "portafolio + pitch), que el docente reparte entre ambos; cada uno se evalúa con su "
    "rúbrica por criterio (bandas Excelente / Aceptable / Insuficiente) en sesion-05.md y "
    "sesion-10.md. Los 10 puntos de "
    "Tareas / investigaciones y los 20 de Laboratorios se distribuyen de forma equitativa "
    "entre las entregas de cada categoría; cada entrega se califica con una rúbrica breve de "
    "tres bandas (Excelente / Aceptable / Insuficiente).",
    "Las herramientas utilizadas (Google AI Studio y Firebase) son gratuitas y se usan "
    "desde el navegador; solo se requiere una cuenta de Google. Para probar la app nativa "
    "en un teléfono físico basta un cable USB (AI Studio la instala desde Chrome).",
]:
    last = clone_after(last, t)

# --- Tabla evaluación (Tabla 1): sin cambios (rubros y puntos iguales) ---

# --- Cronograma (Tabla 2) ---
cron = d.tables[2]

def set_row(r, fecha, tema, alumno, docente):
    cells = cron.rows[r].cells
    set_cell_text(cells[0], fecha)
    set_cell_text(cells[1], tema)
    set_cell_text(cells[2], alumno)
    set_cell_text(cells[3], docente)

def set_exam_row(r, fecha, texto):
    cells = cron.rows[r].cells
    set_cell_text(cells[0], fecha)
    set_cell_text(cells[1], texto)  # celda fusionada

set_row(1, "Semana 1 (6 de julio)",
    "Introducción, UX/UI móvil y diseño con IA",
    "Tarea 1: Investigación sobre principios de UI/UX móvil\n"
    "Laboratorio 1: Primera app generada en AI Studio (prompt / remix)",
    "Exposición sobre:\n- Introducción al curso y al desarrollo de software asistido por IA\n"
    "- Panorama de herramientas (AI Studio Build para PWA y app nativa)\n"
    "- Principios de UI/UX móvil\n- Configuración de cuentas y entorno")

set_row(2, "Semana 2 (13 de julio)",
    "De idea a interfaz y prototipo",
    "Tarea 2: Definir la identidad visual de su app\n"
    "Laboratorio 2: Prototipo navegable con identidad propia",
    "Exposición sobre:\n- Prompting de interfaces\n"
    "- Identidad visual y AI Chips (Nano Banana)\n- Diseño mobile-first\n"
    "- Prototipos interactivos (navegación, estados, microinteracciones)")

set_row(3, "Semana 3 (20 de julio)",
    "Datos y app instalable (PWA)",
    "Tarea 3: Definir el modelo de datos de su app\n"
    "Laboratorio 3: PWA con datos persistentes, instalable",
    "Exposición sobre:\n- Firebase Firestore (auto-provisioning)\n"
    "- Operaciones CRUD y tiempo real (onSnapshot)\n"
    "- Cómo hacer una PWA instalable en el teléfono")

set_row(4, "Semana 4 (27 de julio)",
    "Desarrollo nativo I: fundamentos con IA",
    "Tarea 4: Investigación PWA vs. nativo\n"
    "Laboratorio 4: App nativa creada en AI Studio, vista en el emulador del navegador "
    "(e instalable en el teléfono)",
    "Exposición sobre:\n- Por qué y cuándo conviene nativo\n"
    "- AI Studio para apps nativas (Kotlin + Jetpack Compose): emulador del navegador e 'Install on Device'\n"
    "- Anatomía de una app Android y Material 3 (componentes como vocabulario de diseño)\n"
    "- Cómo llevar la identidad y la UX a nativo")

set_exam_row(5, "Semana 5 (3 de agosto)",
    "Examen Parcial escrito en línea (sesiones 1-4): teoría + práctica escrita. Entrega en línea del Proyecto 1 (PWA instalable con datos)")

set_row(6, "Semana 6 (10 de agosto)",
    "Desarrollo nativo II: login y datos por usuario",
    "Investigación 1: Reglas de seguridad de Firebase\n"
    "Laboratorio 5: App nativa con login con Google y datos privados por usuario en Firestore/Storage",
    "Exposición sobre:\n- Inicio de sesión con Google (Firebase Authentication)\n"
    "- Datos privados por usuario en Firestore (uid, tiempo real)\n"
    "- Fotos en Firebase Storage\n- Depuración de UI con el modo de anotación y capturas de la vista previa")

set_row(7, "Semana 7 (17 de agosto)",
    "Desarrollo nativo III: cámara y pulido de marca",
    "Laboratorio 6: App nativa con cámara, ícono/splash y lista para entregar",
    "Exposición sobre:\n- La cámara del teléfono y permisos en tiempo de ejecución\n"
    "- Reglas de seguridad de Firestore y Storage\n- Ícono, splash y pulido de marca\n"
    "- Depurar con evidencia: capturas de la vista previa y modo de anotación")

set_row(8, "Semana 8 (24 de agosto)",
    "Distribución y marketing de apps",
    "Tarea 5: Estrategia de distribución y marketing de su app\n"
    "Mini-taller: ficha de tienda (store listing)",
    "Exposición / charla (posible invitado) sobre:\n- PWA vs. Play Store\n"
    "- Store listing y ASO básico\n- Monetización de apps\n- Casos de éxito")

set_row(9, "Semana 9 (31 de agosto)",
    "Portafolio, pitch y empleabilidad",
    "Mini-taller: armar el caso de portafolio y el pitch del proyecto",
    "Exposición / charla (posible invitado) sobre:\n"
    "- Portafolio del diseñador-creador con IA\n- Cómo presentar un pitch\n"
    "- Ética y accesibilidad\n- Empleabilidad en el mercado actual")

set_exam_row(10, "Semana 10 (7 de septiembre)",
    "Examen Final escrito en línea (sesiones 5-9): teoría + práctica escrita. Demo Day: showcase y entrega en línea del Proyecto Final (app nativa) + portafolio + pitch")

# --- Ensanchar la columna "Tema" para legibilidad (quitando ancho a "Docente") ---
NEW_W = [1600, 2300, 2200, 2140]  # twips; total = 8240 (igual al original)
grid = cron._tbl.find(qn('w:tblGrid'))
for col, w in zip(grid.findall(qn('w:gridCol')), NEW_W):
    col.set(qn('w:w'), str(w))
for row in cron.rows:
    seen = set()
    for ci, cell in enumerate(row.cells):
        tc = cell._tc
        if id(tc) in seen:   # celda fusionada ya ajustada
            continue
        seen.add(id(tc))
        tcPr = tc.find(qn('w:tcPr'))
        tcW = tcPr.find(qn('w:tcW')) if tcPr is not None else None
        gridspan = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
        span = int(gridspan.get(qn('w:val'))) if gridspan is not None else 1
        width = sum(NEW_W[ci:ci+span])
        if tcW is not None:
            tcW.set(qn('w:w'), str(width))
            tcW.set(qn('w:type'), 'dxa')

d.save(OUT)
print("Guardado:", OUT)
